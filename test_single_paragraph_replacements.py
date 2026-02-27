"""Test replacement of multiple fields in a single paragraph."""

import sys
import os
import json
import tempfile
from docx import Document

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

if 'app' in sys.modules:
    del sys.modules['app']
if 'app.views' in sys.modules:
    del sys.modules['app.views']

from app.views import _populate_template_docx

def test_single_para_multiple_fields():
    """Test multiple fields in a single paragraph."""
    print("=" * 70)
    print("TEST: Multiple Fields in Single Paragraph")
    print("=" * 70)
    
    # Create a document with all three fields in one paragraph
    test_content = """Report for: ACME Corp
Company: ACME Corp
Address: 123 Main St
Company Website: www.acmecorp.com
Contact Company: ACME Corp"""
    
    with tempfile.TemporaryDirectory() as tmpdir:
        template_path = os.path.join(tmpdir, "template.docx")
        doc = Document()
        doc.add_paragraph(test_content)
        doc.save(template_path)
        
        print(f"\nOriginal Template (single paragraph):\n{test_content}\n")
        
        # Define mappings
        mappings = [
            {
                "original": "ACME Corp",
                "field": "company_name",
                "leftContext": "Report for: ",
                "rightContext": "\nCompany:",
                "occurrenceIndex": 1
            },
            {
                "original": "ACME Corp",
                "field": "company_name_2",
                "leftContext": "Company: ",
                "rightContext": "\nAddress:",
                "occurrenceIndex": 1
            },
            {
                "original": "ACME Corp",
                "field": "company_name_3",
                "leftContext": "Contact ",
                "rightContext": "",
                "occurrenceIndex": 1
            }
        ]
        
        placeholders = [m["field"] for m in mappings]
        
        print("Generating augmented document...\n")
        output_doc_path = os.path.join(tmpdir, f"output.docx")
        values = _populate_template_docx(template_path, placeholders, output_doc_path, mappings)
        
        doc = Document(output_doc_path)
        generated_text = "\n".join([p.text for p in doc.paragraphs])
        
        print(f"Generated values: {json.dumps(values, indent=2)}")
        print(f"\nGenerated content:\n{generated_text}\n")
        
        # Check what was replaced
        print("Verification:")
        if f"Report for: {values['company_name']}" in generated_text:
            print(f"✓ company_name correctly replaced: 'Report for: {values['company_name']}'")
        else:
            print(f"✗ company_name NOT replaced in 'Report for: ACME Corp'")
            if "Report for: ACME Corp" in generated_text:
                print(f"  Found: 'Report for: ACME Corp'")
        
        if f"Company: {values['company_name_2']}" in generated_text:
            print(f"✓ company_name_2 correctly replaced: 'Company: {values['company_name_2']}'")
        else:
            print(f"✗ company_name_2 NOT replaced in 'Company: ACME Corp'")
            if "Company: ACME Corp" in generated_text:
                print(f"  Found: 'Company: ACME Corp'")
        
        if f"Contact {values['company_name_3']}" in generated_text:
            print(f"✓ company_name_3 correctly replaced: 'Contact {values['company_name_3']}'")
        else:
            print(f"✗ company_name_3 NOT replaced in 'Contact Company: ACME Corp'")
            if "Contact Company: ACME Corp" in generated_text:
                print(f"  Found: 'Contact Company: ACME Corp'")

if __name__ == "__main__":
    test_single_para_multiple_fields()

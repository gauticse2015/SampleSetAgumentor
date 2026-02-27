"""
Test script for DOCX field replacement with context-based positioning.
Tests that multiple fields are replaced correctly even when in different positions.
"""

import sys
import os
import json
import tempfile
from docx import Document

# Add app to path
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

# Force reload of modules
if 'app' in sys.modules:
    del sys.modules['app']
if 'app.views' in sys.modules:
    del sys.modules['app.views']

from app.views import _populate_template_docx, _generate_value

def create_test_document(content, output_path):
    """Create a test DOCX with the given content."""
    doc = Document()
    doc.add_paragraph(content)
    doc.save(output_path)

def test_multiple_field_replacement():
    """Test that multiple fields are replaced at correct positions."""
    print("=" * 70)
    print("TEST: Multiple Field Replacement with Context")
    print("=" * 70)
    
    # Create test document with multiple fields
    test_content = """Employee Information:
Name: John Doe
Email: john.doe@example.com
Phone: 555-1234
Registration ID: ABC123456
Address: 123 Main Street
City: New York
Amount: 1000.00"""
    
    with tempfile.TemporaryDirectory() as tmpdir:
        template_path = os.path.join(tmpdir, "template.docx")
        output_path = os.path.join(tmpdir, "output.docx")
        
        # Create template
        create_test_document(test_content, template_path)
        print(f"\nOriginal Template:\n{test_content}\n")
        
        # Define field mappings with context
        mappings = [
            {
                "original": "John Doe",
                "field": "full_name",
                "leftContext": "Name: ",
                "rightContext": "\nEmail:",
                "occurrenceIndex": 1
            },
            {
                "original": "john.doe@example.com",
                "field": "email",
                "leftContext": "Email: ",
                "rightContext": "\nPhone:",
                "occurrenceIndex": 1
            },
            {
                "original": "555-1234",
                "field": "phone",
                "leftContext": "Phone: ",
                "rightContext": "\nRegistration",
                "occurrenceIndex": 1
            },
            {
                "original": "ABC123456",
                "field": "registration_id",
                "leftContext": "Registration ID: ",
                "rightContext": "\nAddress:",
                "occurrenceIndex": 1
            },
            {
                "original": "123 Main Street",
                "field": "address",
                "leftContext": "Address: ",
                "rightContext": "\nCity:",
                "occurrenceIndex": 1
            },
            {
                "original": "New York",
                "field": "city",
                "leftContext": "City: ",
                "rightContext": "\nAmount:",
                "occurrenceIndex": 1
            },
            {
                "original": "1000.00",
                "field": "amount",
                "leftContext": "Amount: ",
                "rightContext": "",
                "occurrenceIndex": 1
            }
        ]
        
        # Generate multiple versions
        placeholders = [m["field"] for m in mappings]
        
        print("Generating 3 augmented documents...\n")
        for i in range(3):
            output_doc_path = os.path.join(tmpdir, f"output_{i+1}.docx")
            values = _populate_template_docx(template_path, placeholders, output_doc_path, mappings)
            
            # Read generated document
            doc = Document(output_doc_path)
            generated_text = "\n".join([p.text for p in doc.paragraphs])
            
            print(f"--- Generated Document {i+1} ---")
            print(f"Generated values:\n{json.dumps(values, indent=2)}")
            print(f"\nGenerated content:\n{generated_text}\n")
            
            # Verify all values are present
            all_present = all(str(v) in generated_text for v in values.values())
            if all_present:
                print("✓ All generated values are present in the document")
            else:
                print("✗ FAILED: Some generated values are missing!")
                missing = [k for k, v in values.items() if str(v) not in generated_text]
                print(f"  Missing fields: {missing}")
            print()

def test_identical_fields():
    """Test replacement when same text appears multiple times."""
    print("=" * 70)
    print("TEST: Identical Fields at Different Positions")
    print("=" * 70)
    
    test_content = """Product A - Price: 100
Product B - Price: 200
Product C - Price: 300"""
    
    with tempfile.TemporaryDirectory() as tmpdir:
        template_path = os.path.join(tmpdir, "template.docx")
        output_path = os.path.join(tmpdir, "output.docx")
        
        create_test_document(test_content, template_path)
        print(f"\nOriginal Template:\n{test_content}\n")
        
        # Map each "Price" value separately using context
        mappings = [
            {
                "original": "100",
                "field": "price_a",
                "leftContext": "Product A - Price: ",
                "rightContext": "\nProduct B",
                "occurrenceIndex": 1
            },
            {
                "original": "200",
                "field": "price_b",
                "leftContext": "Product B - Price: ",
                "rightContext": "\nProduct C",
                "occurrenceIndex": 1
            },
            {
                "original": "300",
                "field": "price_c",
                "leftContext": "Product C - Price: ",
                "rightContext": "",
                "occurrenceIndex": 1
            }
        ]
        
        placeholders = [m["field"] for m in mappings]
        
        print("Generating 2 augmented documents...\n")
        for i in range(2):
            output_doc_path = os.path.join(tmpdir, f"output_{i+1}.docx")
            values = _populate_template_docx(template_path, placeholders, output_doc_path, mappings)
            
            doc = Document(output_doc_path)
            generated_text = "\n".join([p.text for p in doc.paragraphs])
            
            print(f"--- Generated Document {i+1} ---")
            print(f"Generated values: {values}")
            print(f"\nGenerated content:\n{generated_text}\n")
            
            # Verify correct placement
            if f"Product A - Price: {values['price_a']}" in generated_text:
                print("✓ Price A correctly placed")
            else:
                print("✗ Price A NOT correctly placed")
            
            if f"Product B - Price: {values['price_b']}" in generated_text:
                print("✓ Price B correctly placed")
            else:
                print("✗ Price B NOT correctly placed")
            
            if f"Product C - Price: {values['price_c']}" in generated_text:
                print("✓ Price C correctly placed")
            else:
                print("✗ Price C NOT correctly placed")
            print()

if __name__ == "__main__":
    test_multiple_field_replacement()
    print("\n\n")
    test_identical_fields()
    print("\n" + "=" * 70)
    print("All tests completed!")
    print("=" * 70)

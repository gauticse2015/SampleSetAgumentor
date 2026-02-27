"""Test with debug logging added to the replacement function."""

import sys
import os
import json
import tempfile
from docx import Document

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

# Manually implement the replacement logic with debug output
def populate_template_docx_debug(doc_path, placeholders, output_path, mappings=None):
    """Debug version of _populate_template_docx."""
    doc = Document(doc_path)
    values = {}
    for field in placeholders:
        # Use simple values for testing
        values[field] = f"VALUE_{field}"
    
    mappings = mappings or []
    
    print(f"Processing document with {len(mappings)} mappings")
    print(f"Placeholders: {placeholders}")
    print(f"Values: {values}\n")
    
    replaced_pairs = set()
    
    def process_paragraph(paragraph, para_id):
        """Process all replacements in a paragraph."""
        print(f"\n=== Processing paragraph {para_id} ===")
        print(f"Original text: {paragraph.text[:80]}...")
        
        replacements = []  # List of (position, original_len, value, mapping_idx)
        text = paragraph.text
        
        # Find all matches for all mappings
        for i in range(len(mappings)):
            pair_key = (i, para_id)
            if pair_key in replaced_pairs:
                print(f"  Mapping {i} already replaced in this paragraph, skipping")
                continue
            
            mapping = mappings[i]
            field = mapping.get('field')
            value = values.get(field)
            
            if not field or not value:
                print(f"  Mapping {i}: no field or value")
                continue
            
            original = mapping.get('original')
            left_context = mapping.get('leftContext', '')
            right_context = mapping.get('rightContext', '')
            occurrence_index = mapping.get('occurrenceIndex', 1)
            
            if not original:
                print(f"  Mapping {i}: no original")
                continue
            
            print(f"  Searching for mapping {i} ({field}): '{original}'")
            
            # Find all occurrences of this mapping in the paragraph
            best_match_idx = -1
            match_count = 0
            search_start = 0
            
            while True:
                idx = text.find(original, search_start)
                if idx == -1:
                    break
                
                print(f"    Found at position {idx}")
                
                # Check context
                context_matches = True
                
                if left_context:
                    lookback = max(len(left_context) * 3, 200)
                    check_start = max(0, idx - lookback)
                    actual_left = text[check_start:idx]
                    if left_context not in actual_left:
                        print(f"      LEFT CONTEXT MISMATCH: '{left_context}' not in '{actual_left[:50]}...'")
                        context_matches = False
                    else:
                        print(f"      LEFT CONTEXT OK")
                
                if right_context and context_matches:
                    lookahead = max(len(right_context) * 3, 200)
                    check_end = min(len(text), idx + len(original) + lookahead)
                    actual_right = text[idx + len(original):check_end]
                    if right_context not in actual_right:
                        print(f"      RIGHT CONTEXT MISMATCH: '{right_context}' not in '...{actual_right[-50:]}'")
                        context_matches = False
                    else:
                        print(f"      RIGHT CONTEXT OK")
                
                if context_matches:
                    match_count += 1
                    print(f"      MATCH #{match_count}")
                    if match_count == occurrence_index:
                        best_match_idx = idx
                        print(f"      THIS IS THE ONE WE WANT!")
                        break
                
                search_start = idx + 1
            
            if best_match_idx != -1:
                replacements.append((best_match_idx, len(original), value, i))
                print(f"    -> Will replace at position {best_match_idx}")
            else:
                print(f"    -> NOT FOUND")
        
        print(f"\n  All replacements for this paragraph: {replacements}")
        
        # Sort replacements by position (descending) so we apply from end to start
        replacements.sort(key=lambda x: x[0], reverse=True)
        print(f"  Sorted (descending): {replacements}")
        
        # Apply replacements from end to start
        for pos, original_len, value, mapping_idx in replacements:
            print(f"  Replacing at {pos}: '{text[pos:pos+original_len]}' -> '{value}'")
            text = text[:pos] + value + text[pos + original_len:]
            replaced_pairs.add((mapping_idx, para_id))
        
        # Update paragraph text if any replacements were made
        if replacements:
            paragraph.text = text
            print(f"  Final text: {paragraph.text[:80]}...")
    
    # First pass: process regular paragraphs
    for para_idx, paragraph in enumerate(doc.paragraphs):
        process_paragraph(paragraph, ('para', para_idx))
    
    doc.save(output_path)
    return values

# Test
test_content = """Report for: ACME Corp
Company: ACME Corp
Address: 123 Main St
Company Website: www.acmecorp.com
Contact Company: ACME Corp"""

with tempfile.TemporaryDirectory() as tmpdir:
    template_path = os.path.join(tmpdir, "template.docx")
    doc = Document()
    doc.add_paragraph(test_content)
    doc.save(template_path)
    
    mappings = [
        {
            "original": "ACME Corp",
            "field": "company_name",
            "leftContext": "Report for: ",
            "rightContext": "\nCompany:",
            "occurrenceIndex": 1
        },
        {
            "original": "ACME Corp",
            "field": "company_name_2",
            "leftContext": "Company: ",
            "rightContext": "\nAddress:",
            "occurrenceIndex": 1
        },
        {
            "original": "ACME Corp",
            "field": "company_name_3",
            "leftContext": "Contact ",
            "rightContext": "",
            "occurrenceIndex": 1
        }
    ]
    
    placeholders = [m["field"] for m in mappings]
    
    output_path = os.path.join(tmpdir, "output.docx")
    values = populate_template_docx_debug(template_path, placeholders, output_path, mappings)
    
    doc = Document(output_path)
    generated_text = "\n".join([p.text for p in doc.paragraphs])
    
    print(f"\n\n=== FINAL RESULT ===")
    print(f"Generated text:\n{generated_text}")

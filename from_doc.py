# pip install python-docx
#    (if you don't have it yet)

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_registration_form(filename="Registration_Form.docx"):
    doc = Document()
    
    # --- Styles ---
    title_style = doc.styles['Heading 1']
    title_style.font.size = Pt(20)
    title_style.font.bold = True
    
    label_style = doc.styles['Normal']
    label_font = label_style.font
    label_font.size = Pt(12)
    label_font.bold = True
    
    normal_style = doc.styles['Normal']
    normal_font = normal_style.font
    normal_font.size = Pt(11)

    # --- Header ---
    title = doc.add_paragraph("REGISTRATION FORM")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.style = 'Heading 1'
    
    doc.add_paragraph()  # spacing
    
    subtitle = doc.add_paragraph(
        "Please fill in your details. Fields marked with * are required.\n"
        "All information will be kept confidential."
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.style = normal_style

    doc.add_paragraph()  # spacing

    # --- Main form content using table ---
    table = doc.add_table(rows=8, cols=2)
    table.style = 'Table Grid'
    table.allow_autofit = False
    table.columns[0].width = Inches(2.2)
    table.columns[1].width = Inches(4.3)

    # Row 0 - Registration ID
    cell = table.cell(0, 0)
    cell.text = "Registration ID"
    run = cell.paragraphs[0].runs[0]
    run.font.bold = True
    run.font.size = Pt(12)

    cell = table.cell(0, 1)
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("REG-____________________")
    run.font.size = Pt(12)
    run.underline = True

    # Row 1 - Full Name
    table.cell(1, 0).text = "Full Name *"
    p = table.cell(1, 1).add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("______________________________________________")
    run.font.size = Pt(12)
    run.underline = True

    # Row 2 - Email
    table.cell(2, 0).text = "Email Address *"
    p = table.cell(2, 1).add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("______________________________________________")
    run.font.size = Pt(12)
    run.underline = True

    # Row 3 - Phone
    table.cell(3, 0).text = "Phone Number"
    p = table.cell(3, 1).add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("______________________________")
    run.font.size = Pt(12)
    run.underline = True

    # Row 4 - Date of Birth
    table.cell(4, 0).text = "Date of Birth"
    p = table.cell(4, 1).add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("DD / MM / YYYY    ________________")
    run.font.size = Pt(12)
    run.underline = True

    # Row 5 - Gender
    table.cell(5, 0).text = "Gender"
    p = table.cell(5, 1).add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("Male / Female / Other / Prefer not to say : ________________________")
    run.font.size = Pt(12)
    run.underline = True

    # Row 6 - Pincode
    table.cell(6, 0).text = "Pincode"
    p = table.cell(6, 1).add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("______")
    run.font.size = Pt(12)
    run.underline = True
    run = p.add_run("   (6 digits)")
    run.font.size = Pt(10)
    run.italic = True

    # Row 7 - Declaration
    table.cell(7, 0).text = "Declaration"
    cell = table.cell(7, 1)
    p = cell.add_paragraph(
        "I hereby declare that the information provided above is true "
        "and correct to the best of my knowledge."
    )
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)

    p = cell.add_paragraph("Signature: ")
    run = p.add_run("___________________________________________")
    run.underline = True
    run.font.size = Pt(12)

    run = p.add_run("          Date: ")
    run.font.size = Pt(12)

    run = p.add_run("________________")
    run.underline = True
    run.font.size = Pt(12)

    # Final spacing and note
    doc.add_paragraph()
    note = doc.add_paragraph("Form generated on " + "February 2025")  # ← you can make dynamic
    note.style = normal_style
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save
    doc.save(filename)
    print(f"File saved as: {filename}")


# ────────────────────────────────────────────────
#   Run the function
# ────────────────────────────────────────────────

if __name__ == "__main__":
    create_registration_form("Registration_Form_Fillable.docx")
"""
Advanced test for field placement including tables and edge cases.
"""

import sys
import os
import json
import tempfile
from docx import Document
from docx.shared import Pt

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

if 'app' in sys.modules:
    del sys.modules['app']
if 'app.views' in sys.modules:
    del sys.modules['app.views']

from app.views import _populate_template_docx

def create_document_with_table(output_path):
    """Create a DOCX with both paragraphs and table."""
    doc = Document()
    
    # Add some intro text
    doc.add_paragraph("Employee Record")
    doc.add_paragraph("Personal Information:")
    
    # Add a table
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Field'
    hdr_cells[1].text = 'Value'
    
    # Data rows
    table.rows[1].cells[0].text = 'Name'
    table.rows[1].cells[1].text = 'John Doe'
    
    table.rows[2].cells[0].text = 'Email'
    table.rows[2].cells[1].text = 'john@example.com'
    
    table.rows[3].cells[0].text = 'Phone'
    table.rows[3].cells[1].text = '555-1234'
    
    # Add closing paragraph
    doc.add_paragraph("End of record")
    
    doc.save(output_path)

def test_table_field_replacement():
    """Test field replacement in tables."""
    print("=" * 70)
    print("TEST: Field Replacement in Tables")
    print("=" * 70)
    
    with tempfile.TemporaryDirectory() as tmpdir:
        template_path = os.path.join(tmpdir, "template_table.docx")
        create_document_with_table(template_path)
        
        # Read and display original
        doc = Document(template_path)
        print("\nOriginal Document:")
        for para in doc.paragraphs:
            if para.text:
                print(f"  {para.text}")
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        print(f"  {cell.text}")
        
        # Define mappings for fields in table
        mappings = [
            {
                "original": "John Doe",
                "field": "full_name",
                "leftContext": "Name",
                "rightContext": "",
                "occurrenceIndex": 1
            },
            {
                "original": "john@example.com",
                "field": "email",
                "leftContext": "Email",
                "rightContext": "",
                "occurrenceIndex": 1
            },
            {
                "original": "555-1234",
                "field": "phone",
                "leftContext": "Phone",
                "rightContext": "",
                "occurrenceIndex": 1
            }
        ]
        
        placeholders = [m["field"] for m in mappings]
        
        print("\n\nGenerating 2 augmented documents with table data...\n")
        for i in range(2):
            output_doc_path = os.path.join(tmpdir, f"output_table_{i+1}.docx")
            values = _populate_template_docx(template_path, placeholders, output_doc_path, mappings)
            
            doc = Document(output_doc_path)
            
            print(f"--- Generated Document {i+1} ---")
            print(f"Generated values: {json.dumps(values, indent=2)}")
            print("\nGenerated content:")
            
            for para in doc.paragraphs:
                if para.text:
                    print(f"  {para.text}")
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            print(f"  {cell.text}")
            
            # Verify values are present
            all_text = "\n".join([p.text for p in doc.paragraphs])
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        all_text += "\n" + cell.text
            
            if all(str(v) in all_text for v in values.values()):
                print("✓ All generated values present in document")
            else:
                print("✗ FAILED: Some values missing")
            print()

def test_repeated_fields():
    """Test replacement of repeated field values."""
    print("=" * 70)
    print("TEST: Repeated Field Values")
    print("=" * 70)
    
    test_content = """Report for: ACME Corp
Company: ACME Corp
Address: 123 Main St
Company Website: www.acmecorp.com
Contact Company: ACME Corp"""
    
    with tempfile.TemporaryDirectory() as tmpdir:
        template_path = os.path.join(tmpdir, "template.docx")
        doc = Document()
        doc.add_paragraph(test_content)
        doc.save(template_path)
        
        print(f"\nOriginal Template:\n{test_content}\n")
        
        # Map the same field name appearing in different contexts
        mappings = [
            {
                "original": "ACME Corp",
                "field": "company_name",
                "leftContext": "Report for: ",
                "rightContext": "\nCompany:",
                "occurrenceIndex": 1
            },
            {
                "original": "ACME Corp",
                "field": "company_name_2",
                "leftContext": "Company: ",
                "rightContext": "\nAddress:",
                "occurrenceIndex": 1
            },
            {
                "original": "ACME Corp",
                "field": "company_name_3",
                "leftContext": "Contact ",
                "rightContext": "",
                "occurrenceIndex": 1
            }
        ]
        
        placeholders = [m["field"] for m in mappings]
        
        print("Generating 1 augmented document...\n")
        output_doc_path = os.path.join(tmpdir, f"output_repeated.docx")
        values = _populate_template_docx(template_path, placeholders, output_doc_path, mappings)
        
        doc = Document(output_doc_path)
        generated_text = "\n".join([p.text for p in doc.paragraphs])
        
        print(f"Generated values: {json.dumps(values, indent=2)}")
        print(f"\nGenerated content:\n{generated_text}\n")
        
        # Verify all replacements worked
        success = True
        if f"Report for: {values['company_name']}" in generated_text:
            print("✓ First occurrence correctly replaced")
        else:
            print("✗ First occurrence NOT replaced")
            success = False
        
        if f"Company: {values['company_name_2']}" in generated_text:
            print("✓ Second occurrence correctly replaced")
        else:
            print("✗ Second occurrence NOT replaced")
            success = False
        
        if f"Contact {values['company_name_3']}" in generated_text:
            print("✓ Third occurrence correctly replaced")
        else:
            print("✗ Third occurrence NOT replaced")
            success = False
        
        print()

if __name__ == "__main__":
    test_table_field_replacement()
    print("\n\n")
    test_repeated_fields()
    print("\n" + "=" * 70)
    print("All advanced tests completed!")
    print("=" * 70)

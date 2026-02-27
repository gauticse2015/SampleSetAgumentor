import os
import zipfile
import io
import time
import csv
import random
import re
import json
import html
from flask import Blueprint, render_template, request, flash, redirect, url_for, current_app, send_from_directory, send_file, session
from werkzeug.utils import secure_filename
from docx import Document
from .utils.augmentor import ImageAugmentor
from .utils.text_augmentor import TextAugmentor

main_bp = Blueprint('main', __name__)

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in current_app.config['ALLOWED_EXTENSIONS']


def _read_docx_text(file_storage):
    document = Document(file_storage)
    paragraphs = [p.text for p in document.paragraphs if p.text]
    return '\n'.join(paragraphs).strip()


def _extract_placeholders(text):
    placeholders = []
    seen = set()
    for match in re.finditer(r"\{\{\s*([\w\-]+)\s*\}\}", text):
        field = match.group(1)
        if field not in seen:
            seen.add(field)
            placeholders.append(field)
    return placeholders


def _generate_value(field_name):
    label = field_name.lower().strip()
    
    # --- Helper functions for random generation ---
    def random_digits(n):
        return ''.join(str(random.randint(0, 9)) for _ in range(n))
        
    def random_letters(n, case='upper'):
        import string
        chars = string.ascii_uppercase if case == 'upper' else string.ascii_letters
        return ''.join(random.choice(chars) for _ in range(n))
        
    def random_alphanumeric(n):
        import string
        chars = string.ascii_uppercase + string.digits
        return ''.join(random.choice(chars) for _ in range(n))

    # --- Field Type Matching ---
    
    # Phone Number (10 digits)
    if any(x in label for x in ['phone', 'mobile', 'cell', 'contact']):
        return f"{random.randint(600, 999)}-{random.randint(100, 999)}-{random.randint(1000, 9999)}"
        
    # Email
    if 'email' in label:
        domains = ['gmail.com', 'yahoo.com', 'outlook.com', 'example.com', 'company.org']
        user = random_letters(random.randint(5, 10), 'lower')
        return f"{user}@{random.choice(domains)}"
        
    # Registration ID / ID (Alphanumeric)
    if any(x in label for x in ['id', 'reg', 'registration', 'reference', 'serial']):
        prefix = random_letters(2)
        nums = random_digits(6)
        return f"{prefix}{nums}"
        
    # Names (Alphabetic only)
    if 'full_name' in label or 'fullname' in label:
        first = random.choice(['James', 'Mary', 'John', 'Patricia', 'Robert', 'Jennifer', 'Michael', 'Linda', 'William', 'Elizabeth'])
        last = random.choice(['Smith', 'Johnson', 'Williams', 'Brown', 'Jones', 'Garcia', 'Miller', 'Davis', 'Rodriguez', 'Martinez'])
        return f"{first} {last}"
        
    if 'first_name' in label or 'firstname' in label or label == 'name':
        return random.choice(['James', 'Mary', 'John', 'Patricia', 'Robert', 'Jennifer', 'Michael', 'Linda', 'William', 'Elizabeth', 'Barbara', 'Richard', 'Susan', 'Joseph', 'Jessica'])
        
    if 'last_name' in label or 'lastname' in label or 'surname' in label:
        return random.choice(['Smith', 'Johnson', 'Williams', 'Brown', 'Jones', 'Garcia', 'Miller', 'Davis', 'Rodriguez', 'Martinez', 'Hernandez', 'Lopez', 'Gonzalez', 'Wilson', 'Anderson'])

    # Address / City / State / Zip
    if 'zip' in label or 'postal' in label or 'pincode' in label:
        return random_digits(5)
        
    if 'city' in label:
        return random.choice(['New York', 'Los Angeles', 'Chicago', 'Houston', 'Phoenix', 'Philadelphia', 'San Antonio', 'San Diego', 'Dallas', 'San Jose'])
        
    if 'state' in label:
        return random.choice(['California', 'Texas', 'Florida', 'New York', 'Pennsylvania', 'Illinois', 'Ohio', 'Georgia', 'North Carolina', 'Michigan'])
        
    if 'country' in label:
        return random.choice(['United States', 'Canada', 'United Kingdom', 'Australia', 'Germany', 'France', 'Japan', 'China', 'India', 'Brazil'])
        
    if 'address' in label:
        num = random.randint(100, 9999)
        street = random.choice(['Main', 'Oak', 'Maple', 'Cedar', 'Elm', 'Washington', 'Lake', 'Hill'])
        st_type = random.choice(['St', 'Ave', 'Rd', 'Blvd', 'Ln', 'Dr'])
        return f"{num} {street} {st_type}"

    # Dates
    if 'date' in label or 'dob' in label or 'birth' in label:
        year = random.randint(1970, 2025)
        month = random.randint(1, 12)
        day = random.randint(1, 28)
        return f"{year}-{month:02d}-{day:02d}"
        
    # Amounts / Money
    if 'amount' in label or 'price' in label or 'cost' in label or 'total' in label:
        return f"{random.randint(10, 5000)}.{random.randint(0, 99):02d}"
        
    # Job Title
    if 'job' in label or 'title' in label or 'role' in label:
        return random.choice(['Engineer', 'Manager', 'Analyst', 'Director', 'Consultant', 'Developer', 'Designer', 'Coordinator', 'Specialist', 'Administrator'])
        
    # Company
    if 'company' in label or 'org' in label:
        return random.choice(['Acme Corp', 'Globex', 'Soylent Corp', 'Initech', 'Umbrella Corp', 'Stark Ind', 'Wayne Ent', 'Cyberdyne', 'Massive Dynamic'])

    # Default fallback: Alphanumeric string
    return f"Val-{random_alphanumeric(6)}"


def _populate_template(text, placeholders):
    populated = text
    values = {}
    for field in placeholders:
        value = _generate_value(field)
        values[field] = value
        populated = re.sub(r"\{\{\s*" + re.escape(field) + r"\s*\}\}", value, populated)
    return populated, values

@main_bp.route('/parse_docx', methods=['POST'])
def parse_docx():
    if 'doc_file' not in request.files:
        return {'error': 'No file part'}, 400
    
    file = request.files['doc_file']
    if file.filename == '':
        return {'error': 'No selected file'}, 400
        
    if file and allowed_file(file.filename):
        try:
            document = Document(file)
            content_parts = []
            for paragraph in document.paragraphs:
                if paragraph.text.strip():
                    content_parts.append(html.escape(paragraph.text))
                    
            # Include table text
            for table in document.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = ' '.join(p.text for p in cell.paragraphs if p.text)
                        row_text.append(cell_text)
                    if row_text:
                        content_parts.append(html.escape(' | '.join(row_text)))
            
            content = '\n'.join(content_parts).strip()
            html_content = '<br>'.join(content_parts)
            return {'content': content, 'html': html_content}
        except Exception as e:
            return {'error': str(e)}, 500
            
    return {'error': 'Invalid file type'}, 400

@main_bp.route('/augment_text', methods=['POST'])
def augment_text():
    text = request.form.get('input_text')
    operations = request.form.getlist('text_operations')
    # ISO 639-3 language code for synonym replacement (default: English)
    lang = request.form.get('synonym_lang', 'eng').strip() or 'eng'

    if not text:
        flash('No text provided')
        return redirect(url_for('main.index'))

    if not operations:
        flash('No text augmentation operations selected')
        return redirect(url_for('main.index'))

    augmentor = TextAugmentor()

    augmented_texts = augmentor.process_text(text, operations, lang=lang)
    
    output_path = current_app.config['OUTPUT_FOLDER']
    if not os.path.exists(output_path):
        os.makedirs(output_path)
        
    timestamp = int(time.time())
    saved_files = []
    
    for i, (aug_type, aug_text) in enumerate(augmented_texts):
        filename = f"text_aug_{timestamp}_{i}_{aug_type.replace(' ', '_').lower()}.txt"
        filepath = os.path.join(output_path, filename)
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(aug_text)
        saved_files.append(filename)
        
    session['generated_text_files'] = saved_files
    flash(f'Successfully generated {len(saved_files)} augmented text files!', 'success')
    return redirect(url_for('main.text_results'))

@main_bp.route('/text_results')
def text_results():
    output_path = current_app.config['OUTPUT_FOLDER']
    files = []
    generated_files = session.get('generated_text_files', [])
    
    if os.path.exists(output_path):
        files = [f for f in generated_files if os.path.exists(os.path.join(output_path, f))]
        
    # Read content for preview
    file_contents = []
    for f in files:
        try:
            with open(os.path.join(output_path, f), 'r') as file:
                content = file.read()
                file_contents.append({'filename': f, 'content': content})
        except:
            pass
            
    return render_template('results.html', text_files=file_contents)


def _populate_template_docx(doc_path, placeholders, output_path, mappings=None):
    doc = Document(doc_path)
    values = {}
    for field in placeholders:
        values[field] = _generate_value(field)
    
    mappings = mappings or []

    # Helper to replace text in runs to preserve formatting
    def replace_text_in_paragraph(paragraph, pending_mappings, values):
        text = paragraph.text
        if not text:
            return

        for i in range(len(pending_mappings) - 1, -1, -1):
            mapping = pending_mappings[i]
            original = mapping.get('original')
            field = mapping.get('field')
            value = values.get(field)
            left_context = mapping.get('leftContext', '')
            right_context = mapping.get('rightContext', '')

            if not original or not value:
                continue

            # Build a context-aware search string if possible
            search_start = None
            if left_context:
                left_index = text.find(left_context)
                if left_index != -1:
                    search_start = left_index + len(left_context)
            else:
                search_start = 0

            # Try to find the original after left_context
            found = text.find(original, search_start)
            if found != -1:
                # Validate right context if provided
                if right_context:
                    after_original = text[found + len(original):]
                    if not after_original.startswith(right_context):
                        continue

                text = text[:found] + value + text[found + len(original):]
                paragraph.text = text
                pending_mappings.pop(i)
                break

    pending_mappings = list(mappings)

    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, pending_mappings, values)
                
    # Handle tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, pending_mappings, values)
                            
    doc.save(output_path)
    return values

@main_bp.route('/augment_dataset', methods=['POST'])
def augment_dataset():
    dataset_count = request.form.get('dataset_count', '0').strip()
    field_mappings_raw = request.form.get('field_mappings', '').strip()
    
    # We need the actual file for DOCX generation to preserve formatting
    doc_file = request.files.get('doc_file')
    has_file = doc_file and doc_file.filename and allowed_file(doc_file.filename)

    if not has_file:
        flash('Please upload a DOCX file before generating the dataset', 'error')
        return redirect(url_for('main.index'))

    try:
        dataset_count = int(dataset_count)
    except ValueError:
        flash('Dataset count must be a number', 'error')
        return redirect(url_for('main.index'))

    if dataset_count < 1:
        flash('Dataset count must be at least 1', 'error')
        return redirect(url_for('main.index'))

    mappings = []
    if field_mappings_raw:
        try:
            mappings = json.loads(field_mappings_raw)
        except json.JSONDecodeError:
            flash('Unable to read highlighted field mappings', 'error')
            return redirect(url_for('main.index'))
    
    # Determine placeholders from highlighted mappings
    placeholders = [mapping.get('field') for mapping in mappings if mapping.get('field')]
    placeholders = [field for field in placeholders if field]

    if not placeholders:
        flash('No highlighted fields found. Please highlight fields in the preview first.', 'error')
        return redirect(url_for('main.index'))

    output_path = current_app.config['OUTPUT_FOLDER']
    os.makedirs(output_path, exist_ok=True)
    
    # Save the uploaded file temporarily if we have it
    temp_doc_path = None
    if has_file:
        temp_filename = f"temp_template_{int(time.time())}.docx"
        temp_doc_path = os.path.join(current_app.config['UPLOAD_FOLDER'], temp_filename)
        doc_file.save(temp_doc_path)

    timestamp = int(time.time())
    csv_filename = f"dataset_{timestamp}.csv"
    csv_path = os.path.join(output_path, csv_filename)

    generated_docs = []
    rows = []
    
    try:
        for i in range(dataset_count):
            if temp_doc_path:
                # Generate DOCX preserving formatting
                doc_filename = f"dataset_{timestamp}_{i + 1}.docx"
                doc_path = os.path.join(output_path, doc_filename)
                values = _populate_template_docx(temp_doc_path, placeholders, doc_path, mappings)
                
                # For preview, we still use text content
                populated_text = f"[DOCX File Generated]\n\nValues used:\n" + "\n".join([f"{k}: {v}" for k,v in values.items()])
                generated_docs.append({'filename': doc_filename, 'content': populated_text})
                rows.append(values)
            else:
                # Fallback to text generation if no file provided (just text area)
                populated_text, values = _populate_template(template_text, placeholders)
                rows.append(values)
                doc_filename = f"dataset_{timestamp}_{i + 1}.txt"
                doc_path = os.path.join(output_path, doc_filename)
                with open(doc_path, 'w', encoding='utf-8') as file:
                    file.write(populated_text)
                generated_docs.append({'filename': doc_filename, 'content': populated_text})

        with open(csv_path, 'w', encoding='utf-8', newline='') as csv_file:
            writer = csv.DictWriter(csv_file, fieldnames=placeholders)
            writer.writeheader()
            writer.writerows(rows)
            
    finally:
        # Clean up temp file
        if temp_doc_path and os.path.exists(temp_doc_path):
            os.remove(temp_doc_path)

    session['generated_dataset_files'] = [doc['filename'] for doc in generated_docs]
    session['generated_dataset_csv'] = csv_filename
    session['generated_dataset_preview'] = generated_docs
    session['generated_dataset_columns'] = placeholders

    flash(f'Successfully generated {dataset_count} dataset records!', 'success')
    return redirect(url_for('main.dataset_results'))


@main_bp.route('/dataset_results')
def dataset_results():
    output_path = current_app.config['OUTPUT_FOLDER']
    preview_docs = session.get('generated_dataset_preview', [])
    csv_filename = session.get('generated_dataset_csv')
    csv_available = csv_filename and os.path.exists(os.path.join(output_path, csv_filename))
    columns = session.get('generated_dataset_columns', [])
    return render_template('results.html', dataset_docs=preview_docs, dataset_csv=csv_filename, dataset_csv_ready=csv_available, dataset_columns=columns)


@main_bp.route('/generated_dataset/<filename>')
def generated_dataset_file(filename):
    return send_from_directory(current_app.config['OUTPUT_FOLDER'], filename)


@main_bp.route('/save_datasets', methods=['POST'])
def save_datasets():
    selected_files = request.form.getlist('selected_dataset_files')
    output_path = request.form.get('output_path', '').strip()

    if not selected_files:
        flash('No dataset files selected for saving', 'error')
        return redirect(url_for('main.dataset_results'))

    if not output_path:
        flash('No output folder path provided', 'error')
        return redirect(url_for('main.dataset_results'))

    output_path = os.path.expanduser(output_path)
    if not os.path.isabs(output_path):
        output_path = os.path.abspath(output_path)

    try:
        os.makedirs(output_path, exist_ok=True)
    except Exception as e:
        flash(f'Error creating output directory: {str(e)}', 'error')
        return redirect(url_for('main.dataset_results'))

    generated_path = current_app.config['OUTPUT_FOLDER']
    saved_count = 0
    errors = []

    for filename in selected_files:
        source_path = os.path.join(generated_path, filename)
        if os.path.exists(source_path):
            try:
                dest_path = os.path.join(output_path, filename)
                counter = 1
                base_name, ext = os.path.splitext(filename)
                while os.path.exists(dest_path):
                    dest_path = os.path.join(output_path, f"{base_name}_{counter}{ext}")
                    counter += 1
                with open(source_path, 'rb') as src, open(dest_path, 'wb') as dst:
                    dst.write(src.read())
                saved_count += 1
            except Exception as e:
                errors.append(f"Error saving {filename}: {str(e)}")
        else:
            errors.append(f"File not found: {filename}")

    if errors:
        for error in errors:
            flash(error, 'error')

    if saved_count > 0:
        flash(f'Successfully saved {saved_count} dataset file(s)', 'success')

    return redirect(url_for('main.dataset_results'))

@main_bp.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        input_mode = request.form.get('input_mode', 'files')
        operations = request.form.getlist('operations')
        
        if not operations:
            flash('No augmentation operations selected')
            return redirect(request.url)
        
        # Process images - append to existing output folder instead of clearing
        output_path = current_app.config['OUTPUT_FOLDER']
        if not os.path.exists(output_path):
            os.makedirs(output_path)
        
        if input_mode == 'folder':
            # Handle folder upload mode
            if 'folder_files[]' not in request.files:
                flash('No folder files part')
                return redirect(request.url)
                
            files = request.files.getlist('folder_files[]')
            
            if not files or files[0].filename == '':
                flash('No files selected in folder')
                return redirect(request.url)

            # Clear previous uploads
            upload_path = current_app.config['UPLOAD_FOLDER']
            if not os.path.exists(upload_path):
                os.makedirs(upload_path)
                
            for f in os.listdir(upload_path):
                os.remove(os.path.join(upload_path, f))
                
            # Save uploaded files
            saved_files = []
            for file in files:
                if file and allowed_file(file.filename):
                    filename = secure_filename(os.path.basename(file.filename))
                    file.save(os.path.join(upload_path, filename))
                    saved_files.append(filename)
                    
            if not saved_files:
                flash('No valid images found in the uploaded folder')
                return redirect(request.url)
            
            # Process images from upload folder (same as single file mode now)
            augmentor = ImageAugmentor(upload_path, output_path)
            
            # Extract configuration from form
            config = {
                'scale_up': request.form.get('scale_up', 20),
                'scale_down': request.form.get('scale_down', 20),
                'brightness_up': request.form.get('brightness_up', 30),
                'brightness_down': request.form.get('brightness_down', 30),
                'contrast_up': request.form.get('contrast_up', 30),
                'contrast_down': request.form.get('contrast_down', 30)
            }
            
            count, generated_files, errors = augmentor.process_images(operations, config)
            
        else:
            # Handle file upload mode
            # Check if the post request has the file part
            if 'files[]' not in request.files:
                flash('No file part')
                return redirect(request.url)
                
            files = request.files.getlist('files[]')
            
            # If user does not select file, browser also
            # submit an empty part without filename
            if not files or files[0].filename == '':
                flash('No selected file')
                return redirect(request.url)

            # Clear previous uploads
            upload_path = current_app.config['UPLOAD_FOLDER']
            if not os.path.exists(upload_path):
                os.makedirs(upload_path)
                
            for f in os.listdir(upload_path):
                os.remove(os.path.join(upload_path, f))
                
            # Save uploaded files
            saved_files = []
            for file in files:
                if file and allowed_file(file.filename):
                    filename = secure_filename(file.filename)
                    file.save(os.path.join(upload_path, filename))
                    saved_files.append(filename)
                    
            if not saved_files:
                flash('No valid images uploaded')
                return redirect(request.url)
            
            augmentor = ImageAugmentor(upload_path, output_path)
            
            # Extract configuration from form
            config = {
                'scale_up': request.form.get('scale_up', 20),
                'scale_down': request.form.get('scale_down', 20),
                'brightness_up': request.form.get('brightness_up', 30),
                'brightness_down': request.form.get('brightness_down', 30),
                'contrast_up': request.form.get('contrast_up', 30),
                'contrast_down': request.form.get('contrast_down', 30)
            }
            
            count, generated_files, errors = augmentor.process_images(operations, config)
        
        # Store generated filenames in session
        session['generated_files'] = generated_files
        
        if errors:
            for error in errors:
                flash(error, 'error')
        
        flash(f'Successfully generated {count} augmented images!', 'success')
        return redirect(url_for('main.results'))

    return render_template('index.html')

@main_bp.route('/results')
def results():
    output_path = current_app.config['OUTPUT_FOLDER']
    images = []
    
    # Only get images that were just generated (from session)
    generated_files = session.get('generated_files', [])
    
    if os.path.exists(output_path):
        # Filter files that exist in output folder and were in the generated list
        images = [f for f in generated_files if os.path.exists(os.path.join(output_path, f))]
        
        # Fallback: if session is empty (e.g. direct access), show nothing or maybe last batch?
        # For now let's stick to showing only what's in session to solve the user's issue
        
    return render_template('results.html', images=images)

@main_bp.route('/generated/<filename>')
def generated_file(filename):
    return send_from_directory(current_app.config['OUTPUT_FOLDER'], filename)

@main_bp.route('/save_texts', methods=['POST'])
def save_texts():
    selected_texts = request.form.getlist('selected_texts')
    output_path = request.form.get('output_path', '').strip()

    if not selected_texts:
        flash('No text files selected for saving', 'error')
        return redirect(url_for('main.text_results'))

    if not output_path:
        flash('No output folder path provided', 'error')
        return redirect(url_for('main.text_results'))

    output_path = os.path.expanduser(output_path)
    if not os.path.isabs(output_path):
        output_path = os.path.abspath(output_path)

    try:
        os.makedirs(output_path, exist_ok=True)
    except Exception as e:
        flash(f'Error creating output directory: {str(e)}', 'error')
        return redirect(url_for('main.text_results'))

    generated_path = current_app.config['OUTPUT_FOLDER']
    saved_count = 0
    errors = []

    for text_filename in selected_texts:
        source_path = os.path.join(generated_path, text_filename)
        if os.path.exists(source_path):
            try:
                dest_path = os.path.join(output_path, text_filename)
                counter = 1
                base_name, ext = os.path.splitext(text_filename)
                while os.path.exists(dest_path):
                    dest_path = os.path.join(output_path, f"{base_name}_{counter}{ext}")
                    counter += 1
                with open(source_path, 'rb') as src, open(dest_path, 'wb') as dst:
                    dst.write(src.read())
                saved_count += 1
            except Exception as e:
                errors.append(f"Error saving {text_filename}: {str(e)}")
        else:
            errors.append(f"File not found: {text_filename}")

    if errors:
        for error in errors:
            flash(error, 'error')

    if saved_count > 0:
        flash(f'Successfully saved {saved_count} text file(s)', 'success')

    return redirect(url_for('main.text_results'))


@main_bp.route('/save', methods=['POST'])
def save_images():
    selected_images = request.form.getlist('selected_images')
    output_path = request.form.get('output_path', '').strip()
    
    if not selected_images:
        flash('No images selected for saving', 'error')
        return redirect(url_for('main.results'))
    
    if not output_path:
        flash('No output folder path provided', 'error')
        return redirect(url_for('main.results'))
    
    # Expand user home directory if path starts with ~
    output_path = os.path.expanduser(output_path)
    
    # Convert to absolute path if relative
    if not os.path.isabs(output_path):
        output_path = os.path.abspath(output_path)
    
    # Create output directory if it doesn't exist
    try:
        os.makedirs(output_path, exist_ok=True)
    except Exception as e:
        flash(f'Error creating output directory: {str(e)}', 'error')
        return redirect(url_for('main.results'))
    
    generated_path = current_app.config['OUTPUT_FOLDER']
    saved_count = 0
    errors = []
    
    for image_name in selected_images:
        source_path = os.path.join(generated_path, image_name)
        if os.path.exists(source_path):
            try:
                dest_path = os.path.join(output_path, image_name)
                # Handle duplicate filenames
                counter = 1
                base_name, ext = os.path.splitext(image_name)
                while os.path.exists(dest_path):
                    dest_path = os.path.join(output_path, f"{base_name}_{counter}{ext}")
                    counter += 1
                
                # Copy file
                with open(source_path, 'rb') as src, open(dest_path, 'wb') as dst:
                    dst.write(src.read())
                saved_count += 1
            except Exception as e:
                errors.append(f"Error saving {image_name}: {str(e)}")
        else:
            errors.append(f"File not found: {image_name}")
    
    if errors:
        for error in errors:
            flash(error, 'error')
    
    if saved_count > 0:
        flash(f'Successfully saved {saved_count} images', 'success')
    
    return redirect(url_for('main.results'))

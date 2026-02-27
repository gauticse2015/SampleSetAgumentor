"""Compare what context is being captured."""

import sys
import os
import json
import tempfile
from docx import Document

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

# Import the actual function
if 'app' in sys.modules:
    del sys.modules['app']
if 'app.views' in sys.modules:
    del sys.modules['app.views']

from app.views import _populate_template_docx

def test_with_actual_context():
    """Test with the actual context captured from frontend."""
    print("=" * 70)
    print("TEST: Using Actual Frontend Context")
    print("=" * 70)
    
    test_content = """Report for: ACME Corp
Company: ACME Corp
Address: 123 Main St
Company Website: www.acmecorp.com
Contact Company: ACME Corp"""
    
    with tempfile.TemporaryDirectory() as tmpdir:
        template_path = os.path.join(tmpdir, "template.docx")
        doc = Document()
        doc.add_paragraph(test_content)
        doc.save(template_path)
        
        # These are the mappings that would come from the frontend
        # after highlighting the fields
        mappings = [
            {
                "original": "ACME Corp",
                "field": "company_name",
                # Frontend captures 100 chars of context
                "leftContext": "Report for: ",
                "rightContext": "\nCompany:",
                "occurrenceIndex": 1,
                "placeholder": "{{ company_name }}"
            },
            {
                "original": "ACME Corp",
                "field": "company_name_2",
                "leftContext": "Company: ",
                "rightContext": "\nAddress:",
                "occurrenceIndex": 1,
                "placeholder": "{{ company_name_2 }}"
            },
            {
                "original": "ACME Corp",
                "field": "company_name_3",
                "leftContext": "Contact ",
                "rightContext": "",
                "occurrenceIndex": 1,
                "placeholder": "{{ company_name_3 }}"
            }
        ]
        
        placeholders = [m["field"] for m in mappings]
        
        print(f"\nOriginal Template:\n{test_content}\n")
        print(f"Mappings being sent:\n{json.dumps(mappings, indent=2)}\n")
        
        output_doc_path = os.path.join(tmpdir, f"output.docx")
        values = _populate_template_docx(template_path, placeholders, output_doc_path, mappings)
        
        doc = Document(output_doc_path)
        generated_text = "\n".join([p.text for p in doc.paragraphs])
        
        print(f"Generated values: {json.dumps(values, indent=2)}")
        print(f"\nGenerated content:\n{generated_text}\n")
        
        # Verify
        print("Verification:")
        if f"Report for: {values['company_name']}" in generated_text:
            print(f"✓ company_name correctly replaced")
        else:
            print(f"✗ company_name NOT replaced")
        
        if f"Company: {values['company_name_2']}" in generated_text:
            print(f"✓ company_name_2 correctly replaced")
        else:
            print(f"✗ company_name_2 NOT replaced")
        
        if f"Contact {values['company_name_3']}" in generated_text:
            print(f"✓ company_name_3 correctly replaced")
        else:
            print(f"✗ company_name_3 NOT replaced")

if __name__ == "__main__":
    test_with_actual_context()
